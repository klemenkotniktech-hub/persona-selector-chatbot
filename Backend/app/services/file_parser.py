import io
import os
import re
import ast
import base64
import json
import traceback
import tempfile
from typing import Tuple, List, Dict, Any, Optional

import pandas as pd
import openpyxl
import docx
from striprtf.striprtf import rtf_to_text
from PIL import Image
from io import BytesIO
import xml.etree.ElementTree as ET
import xmltodict
from langchain_openai import ChatOpenAI
from langchain.schema.messages import HumanMessage
from ..utils.utils import logger

class FileParser:
    """
    Handles parsing of different file types for the chatbot
    """
    
    @staticmethod
    def parse_file(file_content: str, file_name: str, file_type: str) -> Tuple[str, str]:
        """
        Parse file content based on file type
        
        Args:
            file_content: Base64 encoded file content or plain text
            file_name: Name of the file
            file_type: MIME type of the file
            
        Returns:
            Tuple containing (parsed_content, content_type)
            content_type can be 'text', 'table', 'json', 'xml', 'image', or 'binary'
        """
        try:
            # Determine file extension from name
            file_ext = os.path.splitext(file_name)[1].lower()
            
            # Handle different file types
            # Text files
            if file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.java', '.c', '.cpp']:
                return FileParser._parse_text(file_content), 'text'
                
            # PDF files
            elif file_ext == '.pdf':
                return FileParser._parse_pdf(file_content), 'text'
                
            # Word documents
            elif file_ext in ['.docx', '.doc']:
                return FileParser._parse_word(file_content), 'text'
                
            # RTF files
            elif file_ext == '.rtf':
                return FileParser._parse_rtf(file_content), 'text'
                
            # CSV files
            elif file_ext == '.csv':
                return FileParser._parse_csv(file_content), 'table'
                
            # Excel files
            elif file_ext in ['.xlsx', '.xls']:
                return FileParser._parse_excel(file_content), 'table'
                
            # JSON files
            elif file_ext == '.json':
                return FileParser._parse_json(file_content), 'json'
                
            # XML files
            elif file_ext in ['.xml', '.html']:
                return FileParser._parse_xml(file_content), 'xml'
                
            # Image files
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                return FileParser._parse_imported_image(file_content), 'image'
                
            # Default fallback
            else:
                return f"[Unsupported file type: {file_name}]", 'binary'
                
        except Exception as e:
            logger.error(f"Error parsing file {file_name}: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error parsing file {file_name}: {str(e)}]", 'error'
    
    @staticmethod
    def _parse_text(content: str) -> str:
        """Parse plain text content"""
        # If content is base64 encoded, decode it
        if FileParser._is_base64(content):
            try:
                decoded = base64.b64decode(content).decode('utf-8')
                return decoded
            except:
                return content
        return content
    
    @staticmethod
    def _parse_pdf(content) -> str:
        """Parse PDF content using PyMuPDF (fitz)"""
        try:
            logger.info("Starting PDF parsing process with PyMuPDF")
            
            # Check if content is already bytes
            if isinstance(content, bytes):
                logger.info("Content is already in bytes format")
                binary_content = content
            # If content is base64 encoded, decode it
            elif FileParser._is_base64(content):
                logger.info(f"Detected base64 encoded content of length: {len(content)}")
                binary_content = FileParser._extract_base64_content(content)
                logger.info(f"Decoded base64 content to binary of size: {len(binary_content)} bytes")
            else:
                # Assume it's a string that needs to be converted to binary
                logger.info("Content does not appear to be base64 encoded, converting to binary")
                try:
                    # First try UTF-8 encoding
                    binary_content = content.encode('utf-8')
                except Exception as utf8_err:
                    logger.error(f"Error encoding content as UTF-8: {str(utf8_err)}")
                    # If UTF-8 fails, try to treat the content as if it were already binary
                    try:
                        # Try to handle the case where we might have received bytes as a string representation
                        if isinstance(content, str) and content.startswith("b'") and content.endswith("'"):
                            # This might be a string representation of bytes, try to evaluate it
                            logger.info("Content appears to be a string representation of bytes, attempting to parse")
                            try:
                                binary_content = ast.literal_eval(content)
                            except Exception as eval_err:
                                logger.error(f"Error evaluating string as bytes: {str(eval_err)}")
                                binary_content = content.encode('latin-1', errors='ignore')
                        else:
                            binary_content = content.encode('latin-1', errors='ignore')
                    except Exception as latin1_err:
                        logger.error(f"Error encoding content as latin-1: {str(latin1_err)}")
                        # Last resort, try to work with whatever we have
                        if isinstance(content, bytes):
                            binary_content = content
                        else:
                            logger.error("Could not convert content to binary")
                            return "[PDF file could not be processed: Unable to convert content to binary format]"
        
            # Check if we have valid binary content
            if not binary_content:
                logger.error("No binary content extracted from input")
                return "[PDF file could not be processed: No content could be extracted]"
            
            if len(binary_content) < 100:  # A valid PDF should be larger than this
                logger.error(f"Binary content too small to be a valid PDF: {len(binary_content)} bytes")
                return f"[PDF file could not be processed: Content too small ({len(binary_content)} bytes)]"
            
            # Create a file-like object
            pdf_file = io.BytesIO(binary_content)
            
            # Verify this looks like a PDF (should start with %PDF)
            try:
                pdf_start = binary_content[:10].decode('latin-1', errors='ignore')
                logger.info(f"PDF content starts with: {pdf_start}")
                if not pdf_start.startswith('%PDF'):
                    logger.error(f"Content does not appear to be a valid PDF. Starts with: {pdf_start}")
                    # Try to save the binary content to a temporary file for debugging
                    try:
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.bin')
                        temp_file.write(binary_content[:1000])  # Save first 1000 bytes for analysis
                        temp_file.close()
                        logger.info(f"Saved first 1000 bytes of problematic content to {temp_file.name}")
                    except Exception as temp_err:
                        logger.error(f"Error saving debug file: {str(temp_err)}")
                    
                    return "[The file does not appear to be a valid PDF document]"
            except Exception as decode_err:
                logger.error(f"Error checking PDF header: {str(decode_err)}")
            
            # Use PyMuPDF (fitz) to extract text from the PDF
            try:
                # Import fitz (PyMuPDF) here to avoid dependency issues if it's not installed
                import fitz
                
                # Open the PDF with PyMuPDF
                logger.info("Opening PDF with PyMuPDF")
                pdf_file.seek(0)  # Reset file pointer to beginning
                doc = fitz.open(stream=pdf_file, filetype="pdf")
                
                # Get document metadata for logging
                metadata = doc.metadata
                if metadata:
                    title = metadata.get('title', 'Untitled')
                    author = metadata.get('author', 'Unknown')
                    logger.info(f"PDF metadata - Title: {title}, Author: {author}")
                
                # Extract text from each page
                page_count = len(doc)
                logger.info(f"Successfully opened PDF with {page_count} pages")
                
                # Set a reasonable limit for the number of pages to process
                max_pages = min(page_count, 100)  # Process at most 100 pages
                if page_count > max_pages:
                    logger.warning(f"PDF has {page_count} pages, limiting to first {max_pages} pages")
                
                text = ""
                
                # Initialize a global image counter for the entire PDF
                total_image_count = 0
                for page_num in range(max_pages):
                    try:
                        page = doc[page_num]
                        
                        # Get page dimensions for logging
                        width, height = page.rect.width, page.rect.height
                        #logger.info(f"Page {page_num + 1} dimensions: {width}x{height} points")
                        
                        # Extract text with better layout preservation
                        # flags=4 means "preserve layout" which is better for most documents
                        page_text = page.get_text("text", flags=4)
                        
                        # Add page marker and text
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                        
                        # Check if page has images and process them using ChatGPT-4o
                        image_list = page.get_images()
                        if image_list:
                            logger.info(f"Page {page_num + 1} contains {len(image_list)} images")
                            
                            # Extract and process each image
                            for _, img_info in enumerate(image_list):
                                # Increment the global image counter
                                total_image_count += 1
                                
                                try:
                                    xref = img_info[0]  # Get the image reference
                                    base_image = doc.extract_image(xref)
                                    image_bytes = base_image["image"]
                                    
                                    # Convert image bytes to base64 for processing
                                    base64_image = base64.b64encode(image_bytes).decode('utf-8')
                                    
                                    # Use the _parse_image method to analyze the image with ChatGPT-4o
                                    #logger.info(f"Processing image {total_image_count} (on page {page_num + 1}) with ChatGPT-4o")
                                    image_analysis = FileParser._parse_image(base64_image)
                                    
                                    # Add the image analysis to the page text
                                    text += f"\n--- Image {total_image_count} (on Page {page_num + 1}) ---\n{image_analysis}\n"
                                    logger.info(f"Added analysis for image {total_image_count} (on page {page_num + 1})")
                                    
                                except Exception as img_err:
                                    logger.error(f"Error processing image {total_image_count} on page {page_num + 1}: {str(img_err)}")
                                    text += f"\n--- Image {total_image_count} (on Page {page_num + 1}) ---\n[Error analyzing image: {str(img_err)}]\n"
                    
                    except Exception as page_err:
                        logger.error(f"Error extracting text from page {page_num + 1}: {str(page_err)}")
                        text += f"\n--- Page {page_num + 1} ---\n[Error extracting text from this page]"
                
                # Add note if we limited the pages
                if page_count > max_pages:
                    text += f"\n--- Note: Only showing first {max_pages} of {page_count} pages ---\n"
                
                # Close the document
                doc.close()
                
                # Check if we extracted any meaningful text
                if not text.strip():
                    logger.warning("No text content found in PDF")
                    return "[PDF file was processed successfully but no text content was found. The PDF may contain only images or scanned content.]"
                
                # Limit the amount of text returned to avoid overwhelming the system
                max_text_length = 50000  # 50KB of text should be enough for most use cases
                if len(text) > max_text_length:
                    logger.warning(f"Extracted text is very large ({len(text)} chars), truncating to {max_text_length} chars")
                    text = text[:max_text_length] + "\n\n[Text truncated due to size...]\n"
                
                logger.info(f"Successfully extracted {len(text)} characters of text from PDF using PyMuPDF")

                return text
                
            except ImportError as imp_err:
                logger.error(f"PyMuPDF (fitz) is not installed: {str(imp_err)}")
                return "[Error: PyMuPDF is not installed. Please install it with 'pip install pymupdf']"
                
            except Exception as mupdf_err:
                logger.error(f"PyMuPDF error reading PDF: {str(mupdf_err)}")
                return f"[Error reading PDF: {str(mupdf_err)}. The file may be corrupted or password-protected.]"
            
        except MemoryError:
            logger.error("Memory error while processing PDF")
            return "[Error: The PDF file is too large to process]"
        
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error parsing PDF: {str(e)}]"
    
    @staticmethod
    def _parse_csv(content) -> str:
        """Parse CSV content"""
        try:
            # Check if content is already bytes
            if isinstance(content, bytes):
                logger.info("CSV content is already in bytes format")
                # Decode bytes to string for CSV parsing
                try:
                    decoded = content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try other common encodings if UTF-8 fails
                    try:
                        decoded = content.decode('latin-1')
                    except Exception as decode_err:
                        logger.error(f"Error decoding CSV bytes: {str(decode_err)}")
                        return f"[Error decoding CSV: {str(decode_err)}]"
            # If content is base64 encoded, decode it
            elif FileParser._is_base64(content):
                logger.info("CSV content is base64 encoded, decoding")
                binary_content = FileParser._extract_base64_content(content)
                try:
                    decoded = binary_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try other common encodings if UTF-8 fails
                    try:
                        decoded = binary_content.decode('latin-1')
                    except Exception as decode_err:
                        logger.error(f"Error decoding CSV bytes: {str(decode_err)}")
                        return f"[Error decoding CSV: {str(decode_err)}]"
            else:
                # Content is already a string
                logger.info("CSV content is already a string")
                decoded = content
                
            # Parse CSV
            csv_data = io.StringIO(decoded)
            df = pd.read_csv(csv_data)
            
            # Convert to markdown table for better display
            return df.to_markdown(index=False)
        except Exception as e:
            logger.error(f"Error parsing CSV: {str(e)}")
            return f"[Error parsing CSV: {str(e)}]"
    
    @staticmethod
    def _parse_excel(content) -> str:
        """Parse Excel content from all sheets, including images"""
        try:
            # Check if content is already bytes
            if isinstance(content, bytes):
                logger.info("Excel content is already in bytes format")
                binary_content = content
            # If content is base64 encoded, decode it
            elif FileParser._is_base64(content):
                logger.info("Excel content is base64 encoded, decoding")
                binary_content = FileParser._extract_base64_content(content)
            else:
                # Assume it's a string that needs to be converted to binary
                logger.info("Excel content is a string, converting to binary")
                binary_content = content.encode('utf-8')
                
            # Create a file-like object
            excel_file = io.BytesIO(binary_content)
            
            # Get all sheet names using pandas
            xls = pd.ExcelFile(excel_file)
            sheet_names = xls.sheet_names
            
            if not sheet_names:
                return "[Excel file parsed, but no sheets found]" 
            
            # Load the workbook with openpyxl for image extraction
            excel_file.seek(0)
            workbook = openpyxl.load_workbook(excel_file)
            
            # Track total images across all sheets
            total_image_count = 0
            
            # Read all sheets
            all_sheets_content = []
            for sheet_name in sheet_names:
                logger.info(f"Parsing Excel sheet: {sheet_name}")
                
                # Reset file pointer for pandas read
                excel_file.seek(0)
                
                # Read the sheet into a dataframe
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Start building the sheet content
                sheet_content = f"\n### Sheet: {sheet_name}\n"
                
                # Check for images in this sheet using openpyxl
                if sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    
                    # Extract images from this sheet
                    sheet_images = []
                    for image in sheet._images:
                        total_image_count += 1
                        
                        try:
                            # Get image data
                            img_data = image._data()
                            
                            # Convert to base64 for processing with ChatGPT-4o
                            base64_image = base64.b64encode(img_data).decode('utf-8')
                            
                            # Use the _parse_image method to analyze the image
                            logger.info(f"Processing Excel image {total_image_count} from sheet '{sheet_name}'")
                            image_analysis = FileParser._parse_image(base64_image)
                            
                            sheet_images.append(f"\n--- Excel Image {total_image_count} ---\n{image_analysis}")
                            
                        except Exception as img_err:
                            logger.error(f"Error processing Excel image {total_image_count}: {str(img_err)}")
                            sheet_images.append(f"\n--- Excel Image {total_image_count} ---\n[Error analyzing image: {str(img_err)}]")
                    
                    # Add image analyses to sheet content if any were found
                    if sheet_images:
                        sheet_content += f"\n\n**Images found in this sheet:** {len(sheet_images)}\n"
                        sheet_content += "\n".join(sheet_images)
                
                # Add the tabular data
                if df.empty:
                    sheet_content += "\n\n[No tabular data in this sheet]\n"
                else:
                    # Convert to markdown table for better display
                    table = df.to_markdown(index=False)
                    sheet_content += f"\n\n{table}"
                
                all_sheets_content.append(sheet_content)
            
            # Combine all sheets' content
            result = "\n\n".join(all_sheets_content)
            
            # Add summary of images found
            if total_image_count > 0:
                result = f"[Excel file contains {total_image_count} image(s)]\n\n" + result
            
            if not result.strip():
                # If the result is empty, return a message
                return "[Excel file parsed, but no data found or empty spreadsheet]" 
            
            return result
        except Exception as e:
            logger.error(f"Error parsing Excel: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error parsing Excel: {str(e)}]"
    
    @staticmethod
    def _parse_rtf(content) -> str:
        """Parse RTF content"""
        try:
            # Check if content is already bytes
            if isinstance(content, bytes):
                logger.info("RTF content is already in bytes format")
                binary_content = content
            # If content is base64 encoded, decode it
            elif FileParser._is_base64(content):
                logger.info("RTF content is base64 encoded, decoding")
                binary_content = FileParser._extract_base64_content(content)
            else:
                # Assume it's a string that needs to be converted to binary
                logger.info("RTF content is a string, converting to binary")
                binary_content = content.encode('utf-8')
                
            # Convert binary content to string for RTF parsing
            rtf_content = binary_content.decode('latin-1', errors='replace')
            
            # Use striprtf to convert RTF to plain text
            try:
                plain_text = rtf_to_text(rtf_content)
                logger.info(f"Successfully converted RTF to plain text, length: {len(plain_text)}")
            except Exception as rtf_err:
                logger.error(f"Error converting RTF to text: {str(rtf_err)}")
                return f"[Error parsing RTF file: {str(rtf_err)}. The file may be corrupted or in an unsupported format.]"
            
            # Check if we got any meaningful content
            if not plain_text.strip():
                return "[RTF file parsed, but no content found]"
                
            # RTF doesn't support embedded images in the same way as docx,
            # so we'll just return the text content
            return plain_text
            
        except Exception as e:
            logger.error(f"Error parsing RTF file: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error parsing RTF file: {str(e)}]"
    
    @staticmethod
    def _parse_json(content: str) -> str:
        """Parse JSON content"""
        try:
            # If content is base64 encoded, decode it
            if FileParser._is_base64(content):
                decoded = base64.b64decode(content).decode('utf-8')
            else:
                decoded = content
                
            # Parse JSON
            json_data = json.loads(decoded)
            
            # Pretty print JSON
            return json.dumps(json_data, indent=2)
        except Exception as e:
            logger.error(f"Error parsing JSON: {str(e)}")
            return f"[Error parsing JSON: {str(e)}]"
    
    @staticmethod
    def _parse_word(content) -> str:
        """Parse Word document content, including images"""
        try:
            # Check if content is already bytes
            if isinstance(content, bytes):
                binary_content = content
            # If content is string and base64 encoded, decode it
            elif isinstance(content, str) and FileParser._is_base64(content):
                binary_content = FileParser._extract_base64_content(content)
            # If content is string but not base64, encode it
            elif isinstance(content, str):
                binary_content = content.encode('utf-8')
            else:
                # Handle unexpected type
                logger.error(f"Unexpected content type in _parse_word: {type(content)}")
                return f"[Error parsing Word document: Unexpected content type {type(content)}]"
                
            # Create a file-like object
            word_file = io.BytesIO(binary_content)
            
            # Load the document with python-docx
            try:
                doc = docx.Document(word_file)
                logger.info(f"Successfully opened Word document with {len(doc.paragraphs)} paragraphs")
            except Exception as doc_err:
                logger.error(f"Error opening Word document: {str(doc_err)}")
                return f"[Error opening Word document: {str(doc_err)}. The file may be corrupted or in an unsupported format.]"
            
            # Extract text from paragraphs
            text_content = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append(f"\n--- Table {i+1} ---\n" + "\n".join(table_text))
            
            # Extract and process images
            image_count = 0
            image_analyses = []
            
            # Process images from the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_count += 1
                        logger.info(f"Processing Word document image {image_count}")
                        
                        # Get the image data
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        
                        # Convert to base64 for processing with ChatGPT-4o
                        base64_image = base64.b64encode(image_bytes).decode('utf-8')
                        
                        # Use the _parse_image method to analyze the image
                        image_analysis = FileParser._parse_image(base64_image)
                        
                        image_analyses.append(f"\n--- Word Document Image {image_count} ---\n{image_analysis}")
                        
                    except Exception as img_err:
                        logger.error(f"Error processing Word document image {image_count}: {str(img_err)}")
                        image_analyses.append(f"\n--- Word Document Image {image_count} ---\n[Error analyzing image: {str(img_err)}]")
            
            # Combine text content and image analyses
            result = "\n\n".join(text_content)
            
            # Add image analyses if any were found
            if image_analyses:
                result += "\n\n" + "\n".join(image_analyses)
                
            # Add summary of images found
            if image_count > 0:
                result = f"[Word document contains {image_count} image(s)]\n\n" + result
                
            if not result.strip():
                return "[Word document parsed, but no content found]" 
                
            return result
            
        except Exception as e:
            logger.error(f"Error parsing Word document: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Error parsing Word document: {str(e)}]"
    
    @staticmethod
    def _parse_xml(content: str) -> str:
        """Parse XML content"""
        try:
            # If content is base64 encoded, decode it
            if FileParser._is_base64(content):
                decoded = base64.b64decode(content).decode('utf-8')
            else:
                decoded = content
                
            # Parse XML to dict
            xml_dict = xmltodict.parse(decoded)
            
            # Convert to JSON for better readability
            return json.dumps(xml_dict, indent=2)
        except Exception as e:
            logger.error(f"Error parsing XML: {str(e)}")
            return f"[Error parsing XML: {str(e)}]"
    
    @staticmethod
    def _parse_image(content: str) -> str:
        """
        Parse image content using ChatGPT-4o for image description and text extraction
        """
        try:
            # If content is base64 encoded, decode it
            if FileParser._is_base64(content):
                binary_content = FileParser._extract_base64_content(content)
                # Ensure we're working with the base64 string for the API
                # Re-encode to base64 if needed
                base64_image = base64.b64encode(binary_content).decode('utf-8')
            else:
                # Content is already in base64 format
                base64_image = content
            
            # Check if the base64 string has a data URL prefix
            if not base64_image.startswith('data:image'):
                # Add data URL prefix if not present
                base64_image = f"data:image/jpeg;base64,{base64_image}"
        
            logger.info("Sending image to ChatGPT-4o for analysis")
        
            try:
                # Initialize the ChatOpenAI model with vision capabilities
                chat_model = ChatOpenAI(
                    model="gpt-4o",
                    temperature=0.2,  # Lower temperature for more accurate text extraction
                    openai_api_key=os.getenv("OPENAI_API_KEY")
                )
            
                # Create a message with the image
                message = HumanMessage(
                    content=[
                        {"type": "text", "text": "Please describe this image in detail and extract any text visible in it. If there is text, please transcribe it accurately."}, 
                        {"type": "image_url", "image_url": {"url": base64_image}}
                    ]
                )
            
                # Get the response from the model
                response = chat_model.invoke([message])
            
                # Extract the response content
                if hasattr(response, 'content'):
                    analysis_result = response.content
                else:
                    # Fallback for different response formats
                    analysis_result = str(response)
            
                logger.info(f"Received image analysis from ChatGPT-4o: {analysis_result[:100]}...")
            
                return f"[Image Analysis]:\n{analysis_result}\n[End of Image analysis]"
            
            except Exception as api_error:
                logger.error(f"Error using ChatGPT-4o for image analysis: {str(api_error)}")
                logger.error(traceback.format_exc())
            
                return f"[Image attached. Error during processing with ChatGPT-4o: {str(api_error)}. No fallback processing available.]"
        except Exception as e:
            logger.error(f"Error parsing image: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Image attached. Error during processing: {str(e)}]"
    
    @staticmethod
    def _parse_imported_image(content) -> str:
        """
        Parse directly uploaded image files, converting them to base64 first
        and then using the _parse_image method for analysis
        """
        try:
            logger.info("Processing directly uploaded image file")
            
            # Check if content is already bytes
            if isinstance(content, bytes):
                logger.info("Image content is already in bytes format")
                binary_content = content
            # If content is base64 encoded, decode it
            elif FileParser._is_base64(content):
                logger.info("Image content is base64 encoded, decoding")
                binary_content = FileParser._extract_base64_content(content)
            else:
                # Assume it's a string that needs to be converted to binary
                logger.info("Image content is a string, converting to binary")
                binary_content = content.encode('utf-8')
            
            # Convert the binary content to base64 string for _parse_image
            base64_str = base64.b64encode(binary_content).decode('utf-8')
            logger.info(f"Converted image to base64 string of length: {len(base64_str)}")
            
            # Add data URL prefix for image processing
            data_url = f"data:image/jpeg;base64,{base64_str}"
            
            # Use the existing _parse_image method to analyze the image
            return FileParser._parse_image(data_url)
            
        except Exception as e:
            logger.error(f"Error parsing imported image: {str(e)}")
            logger.error(traceback.format_exc())
            return f"[Image attached. Error during processing: {str(e)}]"
    
    @staticmethod
    def _is_base64(content) -> bool:
        """Check if content is likely base64 encoded with improved detection"""
        # Handle bytes input
        if isinstance(content, bytes):
            # Convert first few bytes to string for data URL check
            try:
                # Check if content starts with data URL bytes pattern
                if content.startswith(b'data:'):
                    return True
                    
                # For bytes content, we'll decode a small sample to check if it's valid base64
                # This avoids decoding the entire content which could be large
                sample = content[:1000]  # Take a sample to check
                valid_bytes = set(b'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=')
                valid_byte_count = sum(1 for b in sample if b in valid_bytes)
                return valid_byte_count >= len(sample) * 0.95
            except Exception as e:
                logger.error(f"Error checking bytes in _is_base64: {str(e)}")
                return False
        
        # Handle string input
        # Check if content is a data URL
        if isinstance(content, str) and content.startswith('data:'):
            return True
        
        try:
            # Quick check for non-base64 content
            if len(content) < 8:
                return False
            
            # Clean the content of whitespace and newlines for checking
            cleaned_content = re.sub(r'\s', '', content) if isinstance(content, str) else content
            
            # If the cleaned content is significantly shorter, it likely had a lot of whitespace
            # and is probably not base64
            if isinstance(content, str) and len(cleaned_content) < len(content) * 0.9:
                return False
            
            # Check if most of the content contains base64 characters (allow some tolerance)
            valid_chars = set('ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=')
            valid_char_count = sum(1 for c in cleaned_content if c in valid_chars)
            
            # Consider it base64 if at least 95% of characters are valid base64 characters
            return valid_char_count >= len(cleaned_content) * 0.95
        except Exception as e:
            logger.error(f"Error in _is_base64: {str(e)}")
            return False

    @staticmethod
    def _extract_base64_content(content) -> bytes:
        """Extract binary content from base64 string with improved handling for large files and edge cases"""
        # If content is already bytes, return it directly
        if isinstance(content, bytes):
            logger.info("Content is already in bytes format, returning directly")
            return content
            
        # For debugging purposes, save the content to a file
        try:
            with open("outputBase64.txt", "w", encoding="utf-8") as f:
                f.write(content if isinstance(content, str) else str(content))
        except Exception as e:
            logger.error(f"Error writing debug file: {str(e)}")
        
        try:
            # Determine if content is a data URL or raw base64
            if isinstance(content, str) and content.startswith('data:'):
                logger.info("Detected data URL format")
                # Extract MIME type for logging
                mime_type = "unknown"
                try:
                    mime_part = content.split(',')[0]
                    if ':' in mime_part and ';' in mime_part:
                        mime_type = mime_part.split(':')[1].split(';')[0]
                        logger.info(f"Detected MIME type: {mime_type}")
                except Exception as e:
                    logger.error(f"Error extracting MIME type: {str(e)}")
                
                # Extract the base64 part
                parts = content.split('base64,', 1)
                if len(parts) < 2:
                    logger.error("Invalid data URL format, couldn't find base64 content")
                    return b''
                base64_content = parts[1]
                logger.info(f"Extracted base64 content from data URL, length: {len(base64_content)}")
            else:
                base64_content = content
                logger.info(f"Processing raw base64 content, length: {len(base64_content)}")
            
            # Remove all whitespace first - this is a common issue with base64 data
            base64_content = re.sub(r'\s', '', base64_content)
            logger.info(f"Removed whitespace from base64 content, new length: {len(base64_content)}")
            
            # Check if the content is too large to process in one go
            large_content = len(base64_content) > 1000000  # 1MB threshold
            
            # For large content, use chunked processing
            if large_content:
                logger.info(f"Base64 content is large ({len(base64_content)} chars), using chunked processing")
                chunk_size = 500000  # 500KB chunks
                result = bytearray()
                
                for i in range(0, len(base64_content), chunk_size):
                    chunk = base64_content[i:i+chunk_size]
                    
                    # Clean this chunk (remove invalid characters)
                    chunk = re.sub(r'[^A-Za-z0-9+/]', '', chunk)
                    
                    # Handle padding for each chunk except the last one
                    if i + chunk_size < len(base64_content):
                        # For non-final chunks, make sure length is multiple of 4 by truncating
                        remainder = len(chunk) % 4
                        if remainder > 0:
                            chunk = chunk[:-remainder]
                            logger.info(f"Truncated non-final chunk to ensure multiple of 4 length: {len(chunk)}")
                    else:
                        # For the final chunk, add padding if needed
                        remainder = len(chunk) % 4
                        if remainder > 0:
                            padding_chars = '=' * (4 - remainder)
                            chunk += padding_chars
                            logger.info(f"Added {4 - remainder} padding characters to final chunk: '{padding_chars}'")
                            logger.info(f"Final chunk length after padding: {len(chunk)}")
                            logger.info(f"Last 10 characters of final chunk: '{chunk[-10:]}'")
                    
                    if not chunk:
                        continue
                        
                    try:
                        chunk_bytes = base64.b64decode(chunk)
                        result.extend(chunk_bytes)
                        logger.info(f"Decoded chunk {i//chunk_size + 1}, size: {len(chunk_bytes)} bytes")
                    except Exception as chunk_e:
                        logger.error(f"Error decoding chunk {i//chunk_size + 1}: {str(chunk_e)}")
                        # Continue with next chunk instead of failing completely
                
                if result:
                    logger.info(f"Successfully decoded all chunks, total size: {len(result)} bytes")
                    return bytes(result)
                else:
                    logger.error("Failed to decode any chunks")
                    return b''
            else:
                # For smaller content, process all at once
                # Clean the base64 string (remove invalid characters, including stray padding)
                base64_content = re.sub(r'[^A-Za-z0-9+/]', '', base64_content)
                logger.info(f"Cleaned base64 content (removed all invalid chars and padding), new length: {len(base64_content)}")
                
                # Add padding if needed
                padding_needed = len(base64_content) % 4
                if padding_needed > 0:
                    padding_chars = '=' * (4 - padding_needed)
                    base64_content += padding_chars
                    logger.info(f"Added {4 - padding_needed} padding characters: '{padding_chars}'")
                    logger.info(f"Base64 content length after padding: {len(base64_content)}")
                    logger.info(f"Last 10 characters of content: '{base64_content[-10:]}'")
                
                # Try standard base64 decoding with additional safeguards
                try:
                    # Ensure the length is a multiple of 4 before decoding
                    mod4_len = len(base64_content) % 4
                    if mod4_len != 0:
                        logger.warning(f"Base64 content length ({len(base64_content)}) is not a multiple of 4 (remainder: {mod4_len})")
                        logger.warning(f"This should have been fixed by padding, but checking again")
                        # Force correct padding
                        while len(base64_content) % 4 != 0:
                            base64_content += '='
                        logger.info(f"Re-padded base64 content to length: {len(base64_content)}")
                    
                    # Log some diagnostic information
                    logger.info(f"Attempting to decode base64 content of length {len(base64_content)}")
                    logger.info(f"First 20 chars: '{base64_content[:20]}'")
                    logger.info(f"Last 20 chars: '{base64_content[-20:]}'")                    
                    
                    # Perform the actual decoding
                    result = base64.b64decode(base64_content)
                    logger.info(f"Successfully decoded base64 content, size: {len(result)} bytes")
                    return result
                except Exception as e:
                    logger.error(f"Error decoding base64 content: {str(e)}")
                    logger.error(f"First 100 chars of content: {base64_content[:100]}")
                    logger.error(f"Last 100 chars of content: {base64_content[-100:]}")
                    
                    # Try to identify common base64 issues
                    invalid_chars = [c for c in base64_content if c not in 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=']
                    if invalid_chars:
                        logger.error(f"Found invalid base64 characters: {set(invalid_chars)}")
                        # Try cleaning again
                        base64_content = ''.join(c for c in base64_content if c in 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=')
                        logger.info(f"Re-cleaned base64 content, new length: {len(base64_content)}")
                        
                        # Re-pad if needed
                        while len(base64_content) % 4 != 0:
                            base64_content += '='
                        logger.info(f"Re-padded after cleaning, new length: {len(base64_content)}")
                        
                        # Try decoding again after cleaning
                        try:
                            result = base64.b64decode(base64_content)
                            logger.info(f"Successfully decoded after re-cleaning, size: {len(result)} bytes")
                            return result
                        except Exception as clean_err:
                            logger.error(f"Still failed after re-cleaning: {str(clean_err)}")
                    else:
                        logger.error("No invalid characters found, but decoding still failed")
                    
                    # Try URL-safe base64 decoding as fallback
                    try:
                        logger.info("Attempting URL-safe base64 decoding as fallback")
                        result = base64.urlsafe_b64decode(base64_content)
                        logger.info(f"Successfully decoded using URL-safe base64, size: {len(result)} bytes")
                        return result
                    except Exception as url_safe_err:
                        logger.error(f"URL-safe base64 decoding also failed: {str(url_safe_err)}")
                        
                        # Last resort: try with more lenient padding
                        try:
                            logger.info("Attempting base64 decoding with lenient padding as last resort")
                            # Ensure padding is correct by adding = until length is multiple of 4
                            while len(base64_content) % 4 != 0:
                                base64_content += '='
                            result = base64.b64decode(base64_content, validate=False)
                            logger.info(f"Successfully decoded with lenient padding, size: {len(result)} bytes")
                            return result
                        except Exception as lenient_err:
                            logger.error(f"Lenient base64 decoding also failed: {str(lenient_err)}")
                            return b''
        except MemoryError:
            logger.error("Memory error while processing base64 content (file may be too large)")
            return b''
        except Exception as e:
            logger.error(f"Error extracting base64 content: {str(e)}")
            logger.error(traceback.format_exc())
            return b''